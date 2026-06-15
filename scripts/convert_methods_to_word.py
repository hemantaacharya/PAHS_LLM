"""Convert the publication methods markdown into a Word document."""

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


BASE = Path(__file__).resolve().parent.parent
INPUT_MD = BASE / "METHODS_SECTION_PUBLICATION.md"
OUTPUT_DOC = BASE / "04_results/PAHS_LLM_Methods_2026.docx"


def configure_document(document: Document) -> None:
    """Apply publication-friendly page and font settings."""
    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    heading_settings = {
        "Title": ("Times New Roman", 16, True),
        "Heading 1": ("Times New Roman", 13, True),
        "Heading 2": ("Times New Roman", 11.5, True),
        "Heading 3": ("Times New Roman", 11, True),
    }
    for style_name, (font_name, size, bold) in heading_settings.items():
        style = document.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = bold


def add_paragraph(document: Document, text: str, style: str | None = None, align=None) -> None:
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.add_run(text)


def convert_markdown_to_word() -> None:
    """Convert the markdown methods section into a Word document."""
    if not INPUT_MD.exists():
        raise FileNotFoundError(f"Missing source markdown: {INPUT_MD}")

    document = Document()
    configure_document(document)

    markdown_lines = INPUT_MD.read_text(encoding="utf-8").splitlines()
    title_written = False

    for raw_line in markdown_lines:
        line = raw_line.rstrip()

        if not line.strip():
            continue

        if line.startswith("# ") and not title_written:
            title = document.add_paragraph(style="Title")
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title.paragraph_format.space_after = Pt(10)
            title.add_run(line[2:].strip())
            title_written = True
            continue

        if line.startswith("## "):
            add_paragraph(document, line[3:].strip(), style="Heading 1")
            continue

        if line.startswith("### "):
            add_paragraph(document, line[4:].strip(), style="Heading 2")
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", line)
        if bullet_match:
            add_paragraph(document, bullet_match.group(1).strip(), style="List Bullet")
            continue

        number_match = re.match(r"^\d+\.\s+(.*)$", line)
        if number_match:
            add_paragraph(document, number_match.group(1).strip(), style="List Number")
            continue

        add_paragraph(document, line.strip())

    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOC)
    print(f"Saved methods document to: {OUTPUT_DOC}")


if __name__ == "__main__":
    convert_markdown_to_word()